"""Fill the BCSE497J Project-I report template in place.

The template's formatting IS the design, so nothing here sets a font, size or
spacing. Every paragraph written is a deep copy of a paragraph already in the
template with its text swapped, so it inherits that paragraph's exact
formatting. Section headings are left untouched.

    python tools/build_report.py
"""
from __future__ import annotations

import copy
import sys
from pathlib import Path

import docx
from docx.shared import Inches
from docx.text.paragraph import Paragraph

sys.path.insert(0, str(Path(__file__).resolve().parent))
from report_content import CONTENT, COVER, TOC_ROWS  # noqa: E402

ROOT = Path(__file__).resolve().parents[1]
SRC = Path(r"C:\Users\M V S Akhil Teja\Downloads"
           r"\3 BCSE497J Project I Report - Template.docx")
OUT = ROOT / "report" / "BCSE497J_Project-I_Report_23BCB0135.docx"
FIG = ROOT / "report" / "figures"

# Document order. Everything between one anchor and the next is template
# scaffolding -- instruction lines and /** Sample **/ blocks -- and is removed.
ANCHORS = [
    "ABSTRACT", "TABLE OF CONTENTS", "1. INTRODUCTION", "1.1 Background",
    "1.2 Motivation", "1.3 Scope of the Project",
    "2. PROJECT DESCRIPTION AND GOALS", "2.1 Literature Review",
    "2.2 Research Gap", "2.3 Objectives", "2.4 Problem Statement",
    "2.5 Project Plan", "3. TECHNICAL SPECIFICATION", "3.1 Requirements",
    "3.1.1", "3.1.2", "3.2 Feasibility Study", "3.3 System Specification",
    "3.3.1 Hardware Specification", "3.3.2 Software Specification",
    "4. DESIGN APPROACH AND DETAILS", "4.1 System Architecture", "4.2 Design",
    "4.2.1 Data Flow Diagram", "4.2.2 Use Case Diagram", "5. REFERENCES",
]


def set_text(p: Paragraph, text: str) -> Paragraph:
    """Replace a paragraph's text, keeping the first run's character format."""
    if not p.runs:
        p.add_run(text)
        return p
    p.runs[0].text = text
    for r in list(p.runs[1:]):
        r._element.getparent().remove(r._element)
    return p


def find(doc, prefix, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i >= start and p.text.strip().startswith(prefix):
            return i
    raise KeyError(prefix)


def clone_after(doc, anchor_el, proto_el, text):
    new = copy.deepcopy(proto_el)
    anchor_el.addnext(new)
    set_text(Paragraph(new, doc), text)
    return new


def clear_between(a_el, b_el):
    dead, el = [], a_el.getnext()
    while el is not None and el is not b_el:
        dead.append(el)
        el = el.getnext()
    for el in dead:
        el.getparent().remove(el)
    return len(dead)


def main() -> int:
    doc = docx.Document(str(SRC))

    # Formatting prototypes lifted from the template itself.
    proto_body = copy.deepcopy(
        doc.paragraphs[find(doc, "The rapid advancement of technology")]._element)
    proto_item = copy.deepcopy(
        doc.paragraphs[find(doc, "Data Collection: The system should")]._element)
    proto_h3 = copy.deepcopy(
        doc.paragraphs[find(doc, "3.2.1 Technical Feasibility")]._element)
    proto_cap = copy.deepcopy(
        doc.paragraphs[find(doc, "Fig. 1. Gantt chart")]._element)
    protos = {"body": proto_body, "item": proto_item,
              "h3": proto_h3, "cap": proto_cap}

    # ------------------------------------------------------------ cover ---
    for p in doc.paragraphs[:30]:
        t = p.text.strip()
        if t.startswith("<Title of Project>"):
            set_text(p, COVER["title"])
        elif (t.startswith("(Times New Roman") or t.startswith("(Line spacing")
              or t.startswith("(mention the particulars")):
            set_text(p, "")
        elif t == "(Specialization if any)":
            set_text(p, COVER["specialization"])
        elif t == "September 2026":
            set_text(p, COVER["month_year"])

    t0, t1 = doc.tables[0], doc.tables[1]
    set_text(t0.cell(0, 0).paragraphs[0], COVER["reg"])
    set_text(t0.cell(0, 1).paragraphs[0], COVER["name"])
    for row in list(t0.rows)[1:]:                       # solo project
        row._element.getparent().remove(row._element)
    set_text(t1.cell(0, 0).paragraphs[0], COVER["guide"])
    set_text(t1.cell(1, 0).paragraphs[0], COVER["guide_desig"])

    i_sample = find(doc, "/****** Sample*******/")
    i_abs = find(doc, "ABSTRACT", i_sample)
    n = clear_between(doc.paragraphs[i_sample]._element,
                      doc.paragraphs[i_abs]._element)
    el = doc.paragraphs[i_sample]._element
    el.getparent().remove(el)
    print(f"cover  : filled; removed worked example ({n + 1} elements)")

    # The template numbers two different subsections 3.2. Renumber the second.
    i_sysspec = find(doc, "3.2 System Specification")
    set_text(doc.paragraphs[i_sysspec], "3.3 System Specification")
    for old, new in (("3.2.1 Hardware Specification", "3.3.1 Hardware Specification"),
                     ("3.2.2 Software Specification", "3.3.2 Software Specification")):
        set_text(doc.paragraphs[find(doc, old)], new)

    for p in doc.paragraphs:
        t = p.text
        if "<Mandatory>" in t or "<Optional>" in t:
            set_text(p, t.replace("<Mandatory>", "").replace("<Optional>", "").strip())

    # ---------------------------------------------------- section bodies ---
    anchors, used = {}, 0
    for a in ANCHORS:
        i = find(doc, a, used)
        anchors[a] = doc.paragraphs[i]._element
        used = i + 1

    for a, nxt in zip(ANCHORS, ANCHORS[1:] + [None]):
        a_el = anchors[a]
        if a == "TABLE OF CONTENTS":
            # The contents table sits in this range; clearing it would take
            # the table with it. Its rows are rebuilt further down instead.
            continue
        if nxt is None:
            el = a_el.getnext()
            while el is not None and not el.tag.endswith("}sectPr"):
                nx = el.getnext()
                el.getparent().remove(el)
                el = nx
        else:
            clear_between(a_el, anchors[nxt])

        cur = a_el
        for kind, text in CONTENT.get(a, []):
            if kind == "fig":
                fname, width, caption = text.split("|")
                holder = clone_after(doc, cur, proto_cap, "")
                Paragraph(holder, doc).runs[0].add_picture(
                    str(FIG / fname), width=Inches(float(width)))
                cur = clone_after(doc, holder, proto_cap, caption)
            else:
                cur = clone_after(doc, cur, protos[kind], text)

    # A results section the generic template does not provide, but the review
    # rubric requires. Placed inside section 4 so the numbering stays 1-5.
    i_uc = find(doc, "4.2.2 Use Case Diagram")
    tail = doc.paragraphs[i_uc]._element
    while tail.getnext() is not None and not tail.getnext().tag.endswith("}sectPr"):
        nx = tail.getnext()
        if nx.tag.endswith("}p") and Paragraph(nx, doc).text.strip().startswith("5."):
            break
        tail = nx
    proto_h2 = copy.deepcopy(doc.paragraphs[find(doc, "4.2 Design")]._element)
    cur = clone_after(doc, tail, proto_h2, "4.3 Implementation Status and Results")
    for kind, text in CONTENT["4.3"]:
        cur = clone_after(doc, cur, protos[kind], text)

    # ---------------------------------------------------- table of contents
    toc = doc.tables[-1]
    proto_row = copy.deepcopy(toc.rows[1]._element)
    for row in list(toc.rows)[1:]:
        row._element.getparent().remove(row._element)
    for sl, item, page in TOC_ROWS:
        new = copy.deepcopy(proto_row)
        toc._tbl.append(new)
        cells = toc.rows[-1].cells
        for cell, val in zip(cells, (sl, item, page)):
            set_text(cell.paragraphs[0], val)
            for extra in list(cell.paragraphs[1:]):
                extra._element.getparent().remove(extra._element)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print("wrote  :", OUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
